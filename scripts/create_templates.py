"""Generate placeholder DOCX templates with Jinja2 merge fields.

These templates use docxtpl-compatible {{ variable }} syntax.
All legal content is marked [ATTORNEY: REVIEW AND REPLACE].

Run: python scripts/create_templates.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_BASE = Path(__file__).resolve().parent.parent / "templates"


def add_heading(doc: Document, text: str, level: int = 0) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt({0: 16, 1: 14, 2: 12}.get(level, 12))


def add_para(doc: Document, text: str, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    if align:
        p.alignment = align


def add_centered(doc: Document, text: str, bold: bool = False) -> None:
    add_para(doc, text, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER)


def save(doc: Document, jurisdiction: str, filename: str) -> None:
    out = TEMPLATE_BASE / jurisdiction / filename
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"  Created: {out}")


# ─── DE LLC Templates ─────────────────────────────────────────


def create_certificate_of_formation():
    doc = Document()
    add_centered(doc, "STATE OF DELAWARE", bold=True)
    add_centered(doc, "CERTIFICATE OF FORMATION", bold=True)
    add_centered(doc, "OF", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    add_centered(doc, "A Delaware Limited Liability Company")
    doc.add_paragraph()

    add_para(doc,
        "The undersigned, desiring to form a limited liability company "
        "pursuant to the Delaware Limited Liability Company Act, "
        "6 Del. C. § 18-101 et seq., hereby certifies as follows:"
    )
    doc.add_paragraph()

    add_para(doc, "FIRST: The name of the limited liability company is:", bold=True)
    add_para(doc, "{{ entity_name }}")
    doc.add_paragraph()

    add_para(doc, "SECOND: The address of the registered office in the State of Delaware is:", bold=True)
    add_para(doc, "{{ registered_agent_address }}")
    doc.add_paragraph()

    add_para(doc, "THIRD: The name and address of the registered agent for service of process is:", bold=True)
    add_para(doc, "{{ registered_agent_name }}")
    add_para(doc, "{{ registered_agent_address }}")
    doc.add_paragraph()

    add_para(doc, "IN WITNESS WHEREOF, the undersigned has executed this Certificate of "
        "Formation on {{ formation_date }}.")
    doc.add_paragraph()
    doc.add_paragraph()

    add_para(doc, "By: _______________________________")
    add_para(doc, "Name: {{ responsible_party_name }}")
    add_para(doc, "Title: Authorized Person")

    save(doc, "de_llc", "certificate_of_formation.docx")


def create_operating_agreement():
    doc = Document()
    add_centered(doc, "LIMITED LIABILITY COMPANY AGREEMENT", bold=True)
    add_centered(doc, "OF", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    add_centered(doc, "A Delaware Limited Liability Company")
    add_centered(doc, "Effective as of {{ formation_date }}")
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — This operating agreement template "
        "requires full legal review before use. All substantive provisions "
        "are placeholders.]"
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE I — FORMATION", level=1)
    add_para(doc,
        "1.1 Formation. {{ entity_name }} (the \"Company\") was formed as a "
        "Delaware limited liability company pursuant to the Delaware Limited "
        "Liability Company Act, 6 Del. C. § 18-101, et seq. (the \"Act\") by "
        "filing a Certificate of Formation with the Secretary of State of "
        "Delaware on {{ formation_date }}."
    )
    add_para(doc,
        "1.2 Name. The name of the Company is {{ entity_name }}."
    )
    add_para(doc,
        "1.3 Registered Agent. The registered agent and registered office of "
        "the Company in Delaware is {{ registered_agent_name }}, located at "
        "{{ registered_agent_address }}."
    )
    add_para(doc,
        "1.4 Term. The Company shall continue until dissolved in accordance "
        "with the Act or this Agreement."
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE II — MEMBERS", level=1)
    add_para(doc,
        "2.1 Members. The initial members of the Company are:"
    )

    # Member table using Jinja2 loop
    add_para(doc,
        "{% for member in members %}"
    )
    add_para(doc,
        "  - {{ member.name }} ({{ member.role }}): {{ member.ownership_percentage }}% ownership interest"
    )
    add_para(doc,
        "{% endfor %}"
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE III — MANAGEMENT", level=1)
    add_para(doc,
        "{% if is_manager_managed %}"
    )
    add_para(doc,
        "3.1 Manager-Managed. The Company shall be manager-managed. The "
        "initial managers are:"
    )
    add_para(doc, "{% for mgr in managers %}")
    add_para(doc, "  - {{ mgr.name }}")
    add_para(doc, "{% endfor %}")
    add_para(doc, "{% else %}")
    add_para(doc,
        "3.1 Member-Managed. The Company shall be member-managed. All "
        "members shall have the right to participate in the management "
        "of the Company."
    )
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_heading(doc, "ARTICLE IV — AI AGENT AUTHORITY", level=1)
    add_para(doc,
        "{% if has_agent %}"
    )
    add_para(doc,
        "4.1 Agent Designation. The Company hereby designates "
        "{{ agent_name }} as an authorized AI agent to act on behalf of "
        "the Company within the scope defined in the Agent Authority "
        "Schedule attached hereto as Exhibit A."
    )
    add_para(doc,
        "4.2 Transaction Limits. The agent's per-transaction authorization "
        "limit is ${{ agent_transaction_limit }} USD."
    )
    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Agent authority provisions are novel "
        "and require careful legal review. Consider fiduciary duties, "
        "liability allocation, and regulatory implications.]"
    )
    add_para(doc, "{% else %}")
    add_para(doc,
        "4.1 No AI Agent. The Company has not designated an AI agent at "
        "this time."
    )
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_heading(doc, "ARTICLE V — CAPITAL CONTRIBUTIONS", level=1)
    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Capital contribution provisions.]"
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE VI — DISTRIBUTIONS", level=1)
    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Distribution provisions.]"
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE VII — DISSOLUTION", level=1)
    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Dissolution provisions.]"
    )
    doc.add_paragraph()

    # Signature block
    add_para(doc, "IN WITNESS WHEREOF, the Members have executed this Agreement "
        "as of {{ formation_date }}.", bold=True)
    doc.add_paragraph()

    add_para(doc, "{% for member in members %}")
    add_para(doc, "By: _______________________________")
    add_para(doc, "Name: {{ member.name }}")
    add_para(doc, "")
    add_para(doc, "{% endfor %}")

    save(doc, "de_llc", "operating_agreement.docx")


def create_agent_authority_schedule():
    doc = Document()
    add_centered(doc, "EXHIBIT A", bold=True)
    add_centered(doc, "AGENT AUTHORITY SCHEDULE", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — This Agent Authority Schedule is a "
        "novel legal document defining the scope, limits, and conditions "
        "under which an AI agent may act on behalf of the entity. This "
        "requires careful legal review.]"
    )
    doc.add_paragraph()

    add_heading(doc, "1. Agent Identity", level=1)
    add_para(doc, "Agent Name: {{ agent_name }}")
    add_para(doc, "Agent Version Hash: [TO BE COMPLETED AT DEPLOYMENT]")
    add_para(doc, "{% if agent_smart_contract %}")
    add_para(doc, "Smart Contract Address: {{ agent_smart_contract }}")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_heading(doc, "2. Authority Scope", level=1)
    add_para(doc,
        "The Agent is authorized to perform the following actions on behalf "
        "of {{ entity_name }}:"
    )
    add_para(doc, "{% if agent_authority_scope.get('sign_documents', False) %}")
    add_para(doc, "  - Sign routine business documents")
    add_para(doc, "{% endif %}")
    add_para(doc, "{% if agent_authority_scope.get('manage_compliance', False) %}")
    add_para(doc, "  - Manage compliance filings and deadlines")
    add_para(doc, "{% endif %}")
    add_para(doc, "{% if agent_authority_scope.get('execute_transactions', False) %}")
    add_para(doc, "  - Execute financial transactions")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_heading(doc, "3. Limits", level=1)
    add_para(doc, "Per-transaction limit: ${{ agent_transaction_limit }} USD")
    add_para(doc,
        "The Agent may NOT: [ATTORNEY: REVIEW AND REPLACE — Define "
        "prohibited actions, escalation triggers, and human override "
        "conditions.]"
    )
    doc.add_paragraph()

    add_heading(doc, "4. Reporting", level=1)
    add_para(doc,
        "The Agent shall maintain a complete audit log of all actions "
        "taken on behalf of the Company."
    )
    doc.add_paragraph()

    add_para(doc, "Effective Date: {{ formation_date }}")
    add_para(doc, "Order ID: {{ order_id }}")

    save(doc, "de_llc", "agent_authority_schedule.docx")


def create_form_ss4():
    doc = Document()
    add_centered(doc, "FORM SS-4 — APPLICATION FOR EMPLOYER IDENTIFICATION NUMBER", bold=True)
    add_centered(doc, "[PREPARED FOR IRS SUBMISSION]")
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — This form must be completed "
        "accurately per IRS instructions. The template provides data "
        "mapping only.]"
    )
    doc.add_paragraph()

    # Line items matching SS-4 form
    add_para(doc, "1. Legal name of entity: {{ entity_name }}", bold=True)
    add_para(doc, "2. Trade name (if different): N/A")
    add_para(doc, "3. Executor/administrator/trustee name: N/A")
    add_para(doc, "4a. Mailing address: [TO BE COMPLETED]")
    add_para(doc, "4b. City, state, ZIP: [TO BE COMPLETED]")
    add_para(doc, "5a. Street address: [TO BE COMPLETED]")
    add_para(doc, "5b. City, state, ZIP: [TO BE COMPLETED]")
    add_para(doc, "6. County and state: {{ jurisdiction_full }}")
    add_para(doc, "7a. Name of responsible party: {{ responsible_party_name }}")
    add_para(doc, "7b. SSN/ITIN: [PII — OBTAINED VIA HUMAN KERNEL]")
    add_para(doc, "8a. Is this a LLC?: Yes")
    add_para(doc, "8b. Number of members: {{ member_count }}")
    add_para(doc, "9a. Type of entity: Limited Liability Company")
    add_para(doc, "9b. State of organization: {{ jurisdiction }}")
    add_para(doc, "10. Reason for applying: Started new business")
    add_para(doc, "11. Date business started: {{ formation_date }}")
    add_para(doc, "12. Closing month: December")
    add_para(doc, "13. Highest number of employees expected: 0")
    add_para(doc, "14. Do you expect to pay wages?: No")
    add_para(doc, "15. First date wages paid: N/A")
    add_para(doc, "16. Principal activity: [TO BE COMPLETED]")
    add_para(doc, "17. Principal product/service: [TO BE COMPLETED]")
    add_para(doc, "18. Has the applicant ever applied for an EIN before?: No")
    doc.add_paragraph()

    add_para(doc, "Third Party Designee: {{ responsible_party_name }}")
    add_para(doc, "Designee's telephone: [TO BE COMPLETED]")
    doc.add_paragraph()

    add_para(doc, "Signature: _______________________________")
    add_para(doc, "Name: {{ responsible_party_name }}")
    add_para(doc, "Title: {{ responsible_party_title }}")
    add_para(doc, "Date: {{ formation_date }}")

    save(doc, "de_llc", "form_ss4.docx")


def create_banking_resolution():
    doc = Document()
    add_centered(doc, "BANKING RESOLUTION", bold=True)
    add_centered(doc, "OF", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    doc.add_paragraph()

    add_para(doc,
        "RESOLVED, that {{ entity_name }}, a {{ jurisdiction_full }} "
        "limited liability company (the \"Company\"), hereby authorizes "
        "the opening and maintenance of deposit and operating accounts "
        "at [BANK NAME] (the \"Bank\")."
    )
    doc.add_paragraph()

    add_para(doc,
        "FURTHER RESOLVED, that the following persons are authorized "
        "signers on behalf of the Company:"
    )
    add_para(doc, "{% for member in members %}")
    add_para(doc, "  - {{ member.name }} ({{ member.role }})")
    add_para(doc, "{% endfor %}")
    doc.add_paragraph()

    add_para(doc, "{% if has_agent %}")
    add_para(doc,
        "FURTHER RESOLVED, that the AI agent designated as "
        "{{ agent_name }} is authorized to initiate electronic "
        "transactions on behalf of the Company subject to the limits "
        "set forth in the Agent Authority Schedule, with a per-transaction "
        "limit of ${{ agent_transaction_limit }} USD."
    )
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Banking resolution requires "
        "review to ensure compliance with bank requirements.]"
    )
    doc.add_paragraph()

    add_para(doc, "Effective Date: {{ formation_date }}")
    add_para(doc, "")
    add_para(doc, "{% for member in members %}")
    add_para(doc, "By: _______________________________")
    add_para(doc, "Name: {{ member.name }}")
    add_para(doc, "")
    add_para(doc, "{% endfor %}")

    save(doc, "de_llc", "banking_resolution.docx")


# ─── WY DAO LLC Templates ─────────────────────────────────────


def create_articles_of_organization():
    doc = Document()
    add_centered(doc, "STATE OF WYOMING", bold=True)
    add_centered(doc, "ARTICLES OF ORGANIZATION", bold=True)
    add_centered(doc, "OF", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    add_centered(doc, "A Wyoming Decentralized Autonomous Organization")
    doc.add_paragraph()

    add_para(doc,
        "Pursuant to W.S. § 17-31-104 et seq. (the Wyoming Decentralized "
        "Autonomous Organization Supplement), the undersigned organizer "
        "hereby adopts the following Articles of Organization:"
    )
    doc.add_paragraph()

    add_para(doc, "ARTICLE I — NAME", bold=True)
    add_para(doc, "The name of the organization is {{ entity_name }}.")
    doc.add_paragraph()

    add_para(doc, "ARTICLE II — DAO DESIGNATION", bold=True)
    add_para(doc,
        "This organization is a decentralized autonomous organization "
        "as defined by W.S. § 17-31-102(a)(ix)."
    )
    doc.add_paragraph()

    add_para(doc, "ARTICLE III — REGISTERED AGENT", bold=True)
    add_para(doc, "The registered agent is {{ registered_agent_name }}, "
        "located at {{ registered_agent_address }}.")
    doc.add_paragraph()

    add_para(doc, "ARTICLE IV — SMART CONTRACT", bold=True)
    add_para(doc, "{% if agent_smart_contract %}")
    add_para(doc, "The smart contract identifier for this DAO is: {{ agent_smart_contract }}")
    add_para(doc, "{% else %}")
    add_para(doc, "The smart contract identifier will be designated upon deployment.")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_para(doc, "ARTICLE V — MANAGEMENT", bold=True)
    add_para(doc, "{% if is_manager_managed %}")
    add_para(doc, "The DAO is manager-managed.")
    add_para(doc, "{% else %}")
    add_para(doc, "The DAO is algorithmically managed pursuant to the smart contract.")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Wyoming DAO LLC provisions "
        "require careful review under W.S. § 17-31.]"
    )
    doc.add_paragraph()

    add_para(doc, "Filed on: {{ formation_date }}")
    add_para(doc, "")
    add_para(doc, "By: _______________________________")
    add_para(doc, "Name: {{ responsible_party_name }}")
    add_para(doc, "Title: Organizer")

    save(doc, "wy_dao_llc", "articles_of_organization.docx")


def create_operating_agreement_dao():
    doc = Document()
    add_centered(doc, "OPERATING AGREEMENT", bold=True)
    add_centered(doc, "OF", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    add_centered(doc, "A Wyoming Decentralized Autonomous Organization")
    add_centered(doc, "Effective as of {{ formation_date }}")
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — This DAO operating agreement "
        "must comply with W.S. § 17-31-104 et seq. All provisions are "
        "placeholders pending legal review.]"
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE I — ORGANIZATION", level=1)
    add_para(doc,
        "1.1 Formation. {{ entity_name }} (the \"DAO\") was organized as a "
        "Wyoming DAO LLC pursuant to W.S. § 17-31-104."
    )
    add_para(doc, "1.2 Name. {{ entity_name }}")
    add_para(doc, "1.3 Jurisdiction. Wyoming")
    doc.add_paragraph()

    add_heading(doc, "ARTICLE II — MEMBERS", level=1)
    add_para(doc, "{% for member in members %}")
    add_para(doc, "  - {{ member.name }}: {{ member.ownership_percentage }}%")
    add_para(doc, "{% endfor %}")
    doc.add_paragraph()

    add_heading(doc, "ARTICLE III — SMART CONTRACT GOVERNANCE", level=1)
    add_para(doc,
        "3.1 The DAO's activities shall be governed by its smart contract "
        "as identified in the Smart Contract Identifier Schedule."
    )
    add_para(doc,
        "3.2 In the event of a conflict between the smart contract and "
        "this Agreement, [ATTORNEY: DETERMINE PRIORITY]."
    )
    doc.add_paragraph()

    add_heading(doc, "ARTICLE IV — AI AGENT", level=1)
    add_para(doc, "{% if has_agent %}")
    add_para(doc, "4.1 The DAO designates {{ agent_name }} as its authorized agent.")
    add_para(doc, "{% else %}")
    add_para(doc, "4.1 No AI agent has been designated.")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_para(doc, "[ATTORNEY: REVIEW AND REPLACE — Remaining articles.]")
    doc.add_paragraph()

    add_para(doc, "{% for member in members %}")
    add_para(doc, "By: _______________________________")
    add_para(doc, "Name: {{ member.name }}")
    add_para(doc, "")
    add_para(doc, "{% endfor %}")

    save(doc, "wy_dao_llc", "operating_agreement_dao.docx")


def create_smart_contract_schedule():
    doc = Document()
    add_centered(doc, "SMART CONTRACT IDENTIFIER SCHEDULE", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    doc.add_paragraph()

    add_para(doc,
        "Pursuant to W.S. § 17-31-106(b), the following smart contract(s) "
        "are publicly available and identified for the operation of "
        "{{ entity_name }}:"
    )
    doc.add_paragraph()

    add_para(doc, "Primary Smart Contract:", bold=True)
    add_para(doc, "{% if agent_smart_contract %}")
    add_para(doc, "  Address: {{ agent_smart_contract }}")
    add_para(doc, "  Blockchain: [TO BE SPECIFIED]")
    add_para(doc, "  Deployed: [TO BE COMPLETED]")
    add_para(doc, "{% else %}")
    add_para(doc, "  [SMART CONTRACT NOT YET DEPLOYED]")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_para(doc,
        "[ATTORNEY: REVIEW AND REPLACE — Smart contract schedule must "
        "comply with Wyoming requirements for DAO identification.]"
    )
    doc.add_paragraph()

    add_para(doc, "Effective Date: {{ formation_date }}")
    add_para(doc, "Order ID: {{ order_id }}")

    save(doc, "wy_dao_llc", "smart_contract_schedule.docx")


# ─── Common Templates ─────────────────────────────────────────


def create_bank_pack_cover():
    doc = Document()
    add_centered(doc, "BANK ACCOUNT OPENING PACKAGE", bold=True)
    add_centered(doc, "{{ entity_name }}", bold=True)
    add_centered(doc, "Prepared {{ generated_at }}")
    doc.add_paragraph()

    add_para(doc, "This package contains the following documents for bank "
        "account opening:")
    doc.add_paragraph()

    add_para(doc, "1. Certificate of Formation / Articles of Organization")
    add_para(doc, "2. Operating Agreement")
    add_para(doc, "3. EIN Confirmation Letter")
    add_para(doc, "4. Banking Resolution")
    add_para(doc, "5. Registered Agent Appointment")
    add_para(doc, "{% if has_agent %}")
    add_para(doc, "6. Agent Authority Schedule")
    add_para(doc, "{% endif %}")
    doc.add_paragraph()

    add_para(doc, "Entity Details:", bold=True)
    add_para(doc, "  Name: {{ entity_name }}")
    add_para(doc, "  Type: {{ entity_type }}")
    add_para(doc, "  Jurisdiction: {{ jurisdiction_full }}")
    add_para(doc, "  Formation Date: {{ formation_date }}")
    add_para(doc, "  Order ID: {{ order_id }}")
    doc.add_paragraph()

    add_para(doc, "Beneficial Owners:", bold=True)
    add_para(doc, "{% for member in members %}")
    add_para(doc, "  - {{ member.name }} ({{ member.ownership_percentage }}%)")
    add_para(doc, "{% endfor %}")

    save(doc, "common", "bank_pack_cover.docx")


# ─── Main ──────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Creating DOCX templates...")
    print()

    print("DE LLC templates:")
    create_certificate_of_formation()
    create_operating_agreement()
    create_agent_authority_schedule()
    create_form_ss4()
    create_banking_resolution()

    print()
    print("WY DAO LLC templates:")
    create_articles_of_organization()
    create_operating_agreement_dao()
    create_smart_contract_schedule()

    print()
    print("Common templates:")
    create_bank_pack_cover()

    print()
    print("Done! All templates created.")
